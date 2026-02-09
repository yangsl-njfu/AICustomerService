"""
知识库服务
处理文档上传、文本提取和向量化
"""
import os
import uuid
import json
from typing import List, Dict, Optional, Any
from pathlib import Path
from fastapi import UploadFile
import aiofiles
from pypdf import PdfReader
from docx import Document as DocxDocument
import chardet
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

from config import settings
from .knowledge_retriever import knowledge_retriever


class KnowledgeService:
    """知识库服务类"""

    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR) / "knowledge"
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # 元数据文件，用于存储文档信息
        self.metadata_file = self.upload_dir / "metadata.json"
        self._load_metadata()

        # 文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", "。", "，", " ", ""]
        )
    
    def _load_metadata(self):
        """加载元数据"""
        if self.metadata_file.exists():
            with open(self.metadata_file, 'r', encoding='utf-8') as f:
                self.metadata = json.load(f)
        else:
            self.metadata = {}
    
    def _save_metadata(self):
        """保存元数据"""
        with open(self.metadata_file, 'w', encoding='utf-8') as f:
            json.dump(self.metadata, f, ensure_ascii=False, indent=2)

    def _get_file_extension(self, filename: str) -> str:
        """获取文件扩展名"""
        return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

    def _is_allowed_file(self, filename: str) -> bool:
        """检查文件类型是否允许"""
        ext = self._get_file_extension(filename)
        allowed = ['pdf', 'doc', 'docx', 'txt', 'md']
        return ext in allowed

    async def upload_document(
        self,
        file: UploadFile,
        user_id: str,
        title: Optional[str] = None,
        description: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        上传文档到知识库

        Args:
            file: 上传的文件
            user_id: 上传用户ID
            title: 文档标题（可选）
            description: 文档描述（可选）

        Returns:
            文档信息字典
        """
        try:
            print(f"[DEBUG] 开始上传文档: {file.filename}")
            
            # 检查文件类型
            if not self._is_allowed_file(file.filename):
                raise ValueError(f"不支持的文件类型。允许的类型：pdf, doc, docx, txt, md")

            print(f"[DEBUG] 文件类型检查通过")

            # 检查文件大小
            file.file.seek(0, 2)
            file_size = file.file.tell()
            file.file.seek(0)

            if file_size > settings.MAX_FILE_SIZE:
                raise ValueError(f"文件大小超过限制（最大{settings.MAX_FILE_SIZE / 1024 / 1024}MB）")

            print(f"[DEBUG] 文件大小检查通过: {file_size} bytes")

            # 生成文档ID
            doc_id = str(uuid.uuid4())
            ext = self._get_file_extension(file.filename)
            new_filename = f"{doc_id}.{ext}"
            file_path = self.upload_dir / new_filename

            print(f"[DEBUG] 文档ID: {doc_id}, 保存路径: {file_path}")

            # 保存文件
            content = await file.read()
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(content)

            print(f"[DEBUG] 文件保存成功")

            # 提取文本
            text_content = await self._extract_text(str(file_path), ext)

            print(f"[DEBUG] 文本提取成功，长度: {len(text_content)}")

            if not text_content.strip():
                # 删除空文件
                os.remove(file_path)
                raise ValueError("无法从文件中提取文本内容")

            # 分割文本为 chunks
            chunks = self._split_text(text_content)

            print(f"[DEBUG] 文本分割完成，chunks数量: {len(chunks)}")

            # 准备文档数据
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append({
                    "id": f"{doc_id}_{i}",
                    "content": chunk,
                    "metadata": {
                        "doc_id": doc_id,
                        "doc_title": title or file.filename,
                        "doc_description": description or "",
                        "file_name": file.filename,
                        "file_type": ext,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "uploaded_by": user_id,
                        "source": "knowledge_upload"
                    }
                })

            print(f"[DEBUG] 文档数据准备完成")

            # 添加到向量数据库
            print(f"[DEBUG] knowledge_retriever.available: {knowledge_retriever.available}")
            if knowledge_retriever.available:
                try:
                    print(f"[DEBUG] 开始添加文档到向量数据库")
                    await knowledge_retriever.add_documents(documents, "knowledge_base")
                    print(f"[DEBUG] 文档添加到向量数据库成功")
                except Exception as e:
                    print(f"[WARNING] 向量化失败，但文件已保存: {type(e).__name__}: {str(e)}")
                    # 向量化失败不影响文件上传，继续执行
            else:
                print(f"[WARNING] knowledge_retriever 不可用，跳过向量化")

            print(f"[DEBUG] 上传完成")
            
            # 保存元数据
            self.metadata[doc_id] = {
                "doc_id": doc_id,
                "original_filename": file.filename,
                "title": title or file.filename,
                "description": description or "",
                "file_type": ext,
                "file_size": file_size,
                "chunk_count": len(chunks),
                "uploaded_by": user_id
            }
            self._save_metadata()

            return {
                "doc_id": doc_id,
                "title": title or file.filename,
                "description": description or "",
                "file_name": file.filename,
                "file_type": ext,
                "file_size": file_size,
                "chunk_count": len(chunks),
                "file_path": str(file_path)
            }
        except Exception as e:
            print(f"[ERROR] 上传文档失败: {type(e).__name__}: {str(e)}")
            import traceback
            traceback.print_exc()
            raise

    async def _extract_text(self, file_path: str, ext: str) -> str:
        """从文件提取文本"""
        try:
            if ext == 'pdf':
                return await self._extract_pdf_text(file_path)
            elif ext in ['doc', 'docx']:
                return await self._extract_docx_text(file_path)
            elif ext in ['txt', 'md']:
                return await self._extract_txt_text(file_path)
            else:
                return ""
        except Exception as e:
            print(f"提取文本失败：{e}")
            return ""

    async def _extract_pdf_text(self, file_path: str) -> str:
        """从PDF提取文本"""
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n".join(text_parts)

    async def _extract_docx_text(self, file_path: str) -> str:
        """从Word文档提取文本"""
        doc = DocxDocument(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        return "\n".join(text_parts)

    async def _extract_txt_text(self, file_path: str) -> str:
        """从文本文件提取内容"""
        # 检测编码
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'

        with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
            return f.read()

    def _split_text(self, text: str) -> List[str]:
        """将文本分割成 chunks"""
        # 清理文本
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # 使用 LangChain 的文本分割器
        documents = [LangchainDocument(page_content=text)]
        chunks = self.text_splitter.split_documents(documents)

        return [chunk.page_content for chunk in chunks]

    async def delete_document(self, doc_id: str) -> bool:
        """删除知识库文档"""
        try:
            # 从向量数据库删除
            if knowledge_retriever.available:
                # 获取所有相关 chunks
                # 注意：这里简化处理，实际应该记录所有 chunk IDs
                await knowledge_retriever.delete_document(f"{doc_id}_0", "knowledge_base")

            # 删除文件
            for ext in ['pdf', 'doc', 'docx', 'txt', 'md']:
                file_path = self.upload_dir / f"{doc_id}.{ext}"
                if file_path.exists():
                    os.remove(file_path)
                    break
            
            # 删除元数据
            if doc_id in self.metadata:
                del self.metadata[doc_id]
                self._save_metadata()

            return True
        except Exception as e:
            print(f"删除文档失败：{e}")
            return False

    async def list_documents(self) -> List[Dict[str, Any]]:
        """列出所有知识库文档"""
        try:
            print(f"[DEBUG] 列出文档，目录: {self.upload_dir}")
            documents = []

            if not self.upload_dir.exists():
                print(f"[WARNING] 上传目录不存在: {self.upload_dir}")
                return documents

            for file_path in self.upload_dir.iterdir():
                if file_path.is_file() and file_path.name != "metadata.json":
                    doc_id = file_path.stem
                    stat = file_path.stat()
                    
                    # 从元数据中获取信息
                    if doc_id in self.metadata:
                        meta = self.metadata[doc_id]
                        documents.append({
                            "doc_id": doc_id,
                            "file_name": meta.get("original_filename", file_path.name),
                            "title": meta.get("title", meta.get("original_filename", file_path.name)),
                            "description": meta.get("description", ""),
                            "file_type": file_path.suffix[1:],
                            "file_size": stat.st_size,
                            "chunk_count": meta.get("chunk_count", 0),
                            "created_at": stat.st_ctime
                        })
                    else:
                        # 如果没有元数据，使用文件信息
                        documents.append({
                            "doc_id": doc_id,
                            "file_name": file_path.name,
                            "title": file_path.name,
                            "description": "",
                            "file_type": file_path.suffix[1:],
                            "file_size": stat.st_size,
                            "chunk_count": 0,
                            "created_at": stat.st_ctime
                        })

            print(f"[DEBUG] 找到 {len(documents)} 个文档")
            return documents
        except Exception as e:
            print(f"[ERROR] 列出文档失败: {type(e).__name__}: {str(e)}")
            import traceback
            traceback.print_exc()
            raise


# 全局知识库服务实例
knowledge_service = KnowledgeService()
